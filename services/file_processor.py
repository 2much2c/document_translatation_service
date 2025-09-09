import base64
import io
import PyPDF2
from docx import Document
import pytesseract
from PIL import Image
from typing import Optional
from models.document import Document, FileType
from config.settings import Settings

class FileProcessorService:
    """파일 처리 서비스 클래스"""
    
    def __init__(self):
        """파일 처리 서비스 초기화"""
        self.max_file_size = Settings.MAX_FILE_SIZE
        self.allowed_types = Settings.ALLOWED_FILE_TYPES
    
    def validate_file(self, file_data: bytes, file_type: str, file_name: str) -> bool:
        """파일 유효성 검사"""
        # 파일 크기 검사
        if len(file_data) > self.max_file_size:
            raise ValueError(f"파일 크기가 {self.max_file_size // 1024 // 1024}MB를 초과합니다.")
        
        # 파일 타입 검사
        if file_type not in self.allowed_types:
            raise ValueError(f"지원하지 않는 파일 형식입니다: {file_type}")
        
        return True
    
    def create_document_from_upload(self, file_data: str, file_type: str, file_name: str) -> Document:
        """업로드된 파일로부터 Document 객체 생성"""
        try:
            # Base64 디코딩
            file_bytes = base64.b64decode(file_data)
            
            # 파일 유효성 검사
            self.validate_file(file_bytes, file_type, file_name)
            
            # FileType 열거형으로 변환
            try:
                file_type_enum = FileType(file_type)
            except ValueError:
                raise ValueError(f"지원하지 않는 파일 형식입니다: {file_type}")
            
            # Document 객체 생성
            document = Document(
                file_name=file_name,
                file_type=file_type_enum,
                file_size=len(file_bytes),
                file_data=file_bytes
            )
            
            return document
            
        except Exception as e:
            raise Exception(f"파일 처리 중 오류가 발생했습니다: {str(e)}")
    
    def extract_text_from_document(self, document: Document) -> str:
        """문서에서 텍스트 추출"""
        try:
            file_bytes = document.file_data
            file_type = document.file_type
            
            if file_type == FileType.PDF:
                return self._extract_from_pdf(file_bytes)
            elif file_type in [FileType.DOCX, FileType.DOC]:
                return self._extract_from_docx(file_bytes)
            elif file_type in [FileType.PNG, FileType.JPG, FileType.JPEG, FileType.GIF, FileType.BMP]:
                return self._extract_from_image(file_bytes)
            elif file_type == FileType.TXT:
                return file_bytes.decode('utf-8')
            else:
                raise ValueError(f"지원하지 않는 파일 형식입니다: {file_type}")
                
        except Exception as e:
            raise Exception(f"텍스트 추출 중 오류가 발생했습니다: {str(e)}")
    
    def _extract_from_pdf(self, file_bytes: bytes) -> str:
        """PDF에서 텍스트 추출"""
        try:
            pdf_file = io.BytesIO(file_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"PDF 텍스트 추출 실패: {str(e)}")
    
    def _extract_from_docx(self, file_bytes: bytes) -> str:
        """DOCX에서 텍스트 추출"""
        try:
            doc_file = io.BytesIO(file_bytes)
            doc = Document(doc_file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"DOCX 텍스트 추출 실패: {str(e)}")
    
    def _extract_from_image(self, file_bytes: bytes) -> str:
        """이미지에서 OCR로 텍스트 추출"""
        try:
            image = Image.open(io.BytesIO(file_bytes))
            # OCR 실행
            text = pytesseract.image_to_string(image, lang='kor+eng')
            return text.strip()
        except Exception as e:
            raise Exception(f"이미지 OCR 실패: {str(e)}")
    
    def process_document(self, file_data: str, file_type: str, file_name: str) -> Document:
        """문서 처리 전체 과정"""
        # Document 객체 생성
        document = self.create_document_from_upload(file_data, file_type, file_name)
        
        # 텍스트 추출
        extracted_text = self.extract_text_from_document(document)
        document.extracted_text = extracted_text
        
        return document
